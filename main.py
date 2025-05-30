import tkinter as tk
import threading
import json
import os
import re
import sys
import time
import shutil
import webbrowser
import platform
import subprocess
import logging

# Kritik modüllerin varlığını kontrol et
try:
    import requests
except ImportError:
    print("❌ 'requests' modülü bulunamadı. Güncelleme kontrolü çalışmayacak.")
    requests = None

try:
    from PIL import Image, ImageTk
except ImportError:
    print("❌ 'PIL/Pillow' modülü bulunamadı. Görsel işleme çalışmayacak.")
    sys.exit(1)

try:
    import keyboard
except ImportError:
    print("⚠️ 'keyboard' modülü bulunamadı. Kısayol tuşları çalışmayacak.")
    keyboard = None

from tkinter import ttk, scrolledtext, filedialog, messagebox
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.screen_capture import capture_screen_region
from core.ocr_engine import ocr_image
from core.db_manager import DatabaseManager
from core.gif_capture import record_gif
from core.file_manager import FileManager
from core.system_tray import SystemTray
from core.sound_manager import SoundManager
from core.region_selector import RegionSelector
from core.language import get_text, get_available_languages, get_language_code

# Versiyon bilgileri
CURRENT_VERSION = "2.0.0"
VERSION_CHECK_URL = "https://raw.githubusercontent.com/muallimun/fadim/refs/heads/main/fadim_version.txt"
DOWNLOAD_URL = "https://github.com/muallimun/fadim/releases/latest"
WEBSITE_URL = "https://www.muallimun.com/fadim/"

def check_for_updates():
    """Güncelleme kontrolü yapar"""
    if requests is None:
        print("❌ requests modülü yok, güncelleme kontrolü yapılamıyor")
        return None
    
    try:
        response = requests.get(VERSION_CHECK_URL, timeout=5)
        if response.status_code == 200:
            latest_version = response.text.strip()
            if latest_version != CURRENT_VERSION:
                return latest_version
        return None
    except Exception as e:
        print(f"Güncelleme kontrolü hatası: {e}")
        return None

def show_update_dialog(latest_version, app):
    """Güncelleme dialog'u gösterir"""
    result = messagebox.askyesno(
        app.get_text('update_available_title'),
        app.get_text('update_available_message', current_version=CURRENT_VERSION, latest_version=latest_version)
    )

    if result:
        webbrowser.open(DOWNLOAD_URL)



def check_tesseract_available():
    """Gelişmiş Tesseract kontrol sistemi"""
    if shutil.which("tesseract") is None:
        # Tesseract bulunamadı
        response = messagebox.askyesno(
            "🔧 OCR Programı Gerekli",
            "FADIM'in OCR özelliği için Tesseract OCR programı gereklidir.\n\n" +
            "❌ Tesseract sisteminizde bulunamadı\n\n" +
            "✅ Kurulum için yardım almak ister misiniz?\n\n" +
            "(Hayır seçerseniz, sadece GIF kayıt özelliği kullanılabilir)"
        )

        if response:
            # Yardım penceresini göster
            show_tesseract_installation_guide()
        else:
            # OCR olmadan devam et uyarısı
            messagebox.showinfo(
                get_text('limited_mode'),
                get_text('limited_mode_message')
            )
        return False
    else:
        # Tesseract bulundu, dil paketi kontrolü
        try:
            import subprocess
            result = subprocess.run(['tesseract', '--list-langs'], 
                                  capture_output=True, text=True, timeout=5)
            available_langs = result.stdout.strip().split('\n')[1:]  # İlk satır başlık

            required_langs = ['eng', 'tur', 'ara']
            missing_langs = [lang for lang in required_langs if lang not in available_langs]

            if missing_langs:
                lang_names = {'eng': 'İngilizce', 'tur': 'Türkçe', 'ara': 'Arapça'}
                missing_names = [lang_names[lang] for lang in missing_langs]

                messagebox.showwarning(
                    "⚠️ Eksik Dil Paketleri",
                    f"Tesseract bulundu ancak şu dil paketleri eksik:\n\n" +
                    f"❌ {', '.join(missing_names)}\n\n" +
                    f"Mevcut diller: {', '.join(available_langs)}\n\n" +
                    f"Tam destek için eksik dil paketlerini kurun."
                )
            else:
                print("✅ Tesseract ve tüm dil paketleri mevcut")

        except Exception as e:
            print(f"Dil paketi kontrol hatası: {e}")

        return True

def show_tesseract_installation_guide(current_language='tr'):
    """Tesseract kurulum rehberini göster"""
    import platform
    system = platform.system()

    guide_window = tk.Toplevel()

    # Dil bazında başlık
    if current_language == 'tr':
        guide_window.title("🔧 Tesseract OCR Kurulum Rehberi")
        system_title = f"🖥️ {system} İçin Tesseract Kurulumu"
        steps_title = "📋 Kurulum Adımları"
        close_text = "❌ Kapat"
        download_text = "💾 Windows İndir"
        website_text = "🌐 Resmi Site"
    elif current_language == 'en':
        guide_window.title("🔧 Tesseract OCR Installation Guide")
        system_title = f"🖥️ Tesseract Installation for {system}"
        steps_title = "📋 Installation Steps"
        close_text = "❌ Close"
        download_text = "💾 Download Windows"
        website_text = "🌐 Official Website"
    else:  # Arabic
        guide_window.title("🔧 دليل تثبيت Tesseract OCR")
        system_title = f"🖥️ تثبيت Tesseract لـ {system}"
        steps_title = "📋 خطوات التثبيت"
        close_text = "❌ إغلاق"
        download_text = "💾 تحميل Windows"
        website_text = "🌐 الموقع الرسمي"

    guide_window.geometry("600x500")
    guide_window.resizable(True, True)
    guide_window.attributes('-topmost', True)  # Her zaman üstte

    main_frame = ttk.Frame(guide_window)
    main_frame.pack(fill='both', expand=True, padx=15, pady=15)

    # Başlık
    ttk.Label(main_frame, text=system_title, 
              font=('Arial', 14, 'bold')).pack(pady=(0, 15))

    # Adımlar
    steps_frame = ttk.LabelFrame(main_frame, text=steps_title)
    steps_frame.pack(fill='both', expand=True, pady=(0, 10))

    if system == "Windows":
        if current_language == 'tr':
            steps_text = """1️⃣ İndirme:
   • Aşağıdaki 'Windows İndir' butonuna tıklayın
   • tesseract-ocr-w64-setup-5.3.x.exe dosyasını indirin

2️⃣ Kurulum:
   • İndirilen dosyayı yönetici olarak çalıştırın
   • Kurulum sihirbazında 'Additional script data' seçin
   • Turkish ve Arabic dil paketlerini seçin

3️⃣ PATH Ayarı (ÖNEMLİ!):
   • Windows + R tuşlarına basın, 'sysdm.cpl' yazın
   • 'Gelişmiş' sekmesinde 'Çevre Değişkenleri'ne tıklayın
   • 'Path' değişkenini seçin ve 'Düzenle' tıklayın
   • 'Yeni' ile şu yolu ekleyin: C:\\Program Files\\Tesseract-OCR
   • Tamam ile kaydedin ve bilgisayarı yeniden başlatın

4️⃣ Test:
   • Komut istemi açın (cmd)
   • 'tesseract --version' yazın
   • Sürüm görünürse başarılı!"""
        elif current_language == 'en':
            steps_text = """1️⃣ Download:
   • Click 'Download Windows' button below
   • Download tesseract-ocr-w64-setup-5.3.x.exe file

2️⃣ Installation:
   • Run the downloaded file as administrator
   • Select 'Additional script data' in setup wizard
   • Choose Turkish and Arabic language packs

3️⃣ PATH Setting (IMPORTANT!):
   • Press Windows + R keys, type 'sysdm.cpl'
   • Click 'Environment Variables' in 'Advanced' tab
   • Select 'Path' variable and click 'Edit'
   • Add new path: C:\\Program Files\\Tesseract-OCR
   • Save with OK and restart computer

4️⃣ Test:
   • Open command prompt (cmd)
   • Type 'tesseract --version'
   • If version appears, installation successful!"""
        else:  # Arabic
            steps_text = """1️⃣ التحميل:
   • اضغط على زر 'تحميل Windows' أدناه
   • حمل ملف tesseract-ocr-w64-setup-5.3.x.exe

2️⃣ التثبيت:
   • شغل الملف المحمل كمدير
   • اختر 'Additional script data' في معالج التثبيت
   • اختر حزم اللغة التركية والعربية

3️⃣ إعداد PATH (مهم!):
   • اضغط مفاتيح Windows + R، اكتب 'sysdm.cpl'
   • اضغط 'Environment Variables' في تبويب 'Advanced'
   • اختر متغير 'Path' واضغط 'Edit'
   • أضف مسار جديد: C:\\Program Files\\Tesseract-OCR
   • احفظ بـ OK وأعد تشغيل الكمبيوتر

4️⃣ الاختبار:
   • افتح موجه الأوامر (cmd)
   • اكتب 'tesseract --version'
   • إذا ظهر الإصدار، التثبيت ناجح!"""
    else:  # Linux/Mac
        if current_language == 'tr':
            steps_text = """1️⃣ Terminal Açın:
   • Ctrl+Alt+T ile terminal açın

2️⃣ Kurulum Komutları:
   Ubuntu/Debian:
   sudo apt update
   sudo apt install tesseract-ocr tesseract-ocr-tur tesseract-ocr-ara

   CentOS/RHEL:
   sudo yum install tesseract tesseract-langpack-tur

   Fedora:
   sudo dnf install tesseract tesseract-langpack-tur

3️⃣ Test:
   tesseract --version
   tesseract --list-langs"""
        elif current_language == 'en':
            steps_text = """1️⃣ Open Terminal:
   • Open terminal with Ctrl+Alt+T

2️⃣ Installation Commands:
   Ubuntu/Debian:
   sudo apt update
   sudo apt install tesseract-ocr tesseract-ocr-tur tesseract-ocr-ara

   CentOS/RHEL:
   sudo yum install tesseract tesseract-langpack-tur

   Fedora:
   sudo dnf install tesseract tesseract-langpack-tur

3️⃣ Test:
   tesseract --version
   tesseract --list-langs"""
        else:  # Arabic
            steps_text = """1️⃣ افتح الطرفية:
   • افتح الطرفية بـ Ctrl+Alt+T

2️⃣ أوامر التثبيت:
   Ubuntu/Debian:
   sudo apt update
   sudo apt install tesseract-ocr tesseract-ocr-tur tesseract-ocr-ara

   CentOS/RHEL:
   sudo yum install tesseract tesseract-langpack-tur

   Fedora:
   sudo dnf install tesseract tesseract-langpack-tur

3️⃣ الاختبار:
   tesseract --version
   tesseract --list-langs"""

    steps_label = ttk.Label(steps_frame, text=steps_text, font=('Courier', 9), justify='left')
    steps_label.pack(anchor='w', padx=10, pady=10)

    # Butonlar
    button_frame = ttk.Frame(main_frame)
    button_frame.pack(fill='x', pady=(10, 0))

    if system == "Windows":
        ttk.Button(button_frame, text=download_text, 
                  command=lambda: webbrowser.open("https://digi.bib.uni-mannheim.de/tesseract/")).pack(side='left', padx=5)

    ttk.Button(button_frame, text=website_text, 
              command=lambda: webbrowser.open("https://github.com/tesseract-ocr/tesseract")).pack(side='left', padx=5)

    ttk.Button(button_frame, text=close_text, 
              command=guide_window.destroy).pack(side='right', padx=5)

# Replit ortamı için VNC kontrolü
def setup_replit_display():
    """Replit ortamında VNC display ayarı"""
    if 'REPL_ID' in os.environ:
        print("🖥️ Replit ortamı algılandı, VNC display ayarlanıyor...")
        if 'DISPLAY' not in os.environ:
            os.environ['DISPLAY'] = ':0'
        # X11 auth dosyası yoksa oluştur
        xauth_file = os.path.expanduser('~/.Xauthority')
        if not os.path.exists(xauth_file):
            try:
                open(xauth_file, 'a').close()
                os.chmod(xauth_file, 0o600)
            except:
                pass
        return True
    return False

class FADIM(tk.Tk):
    def __init__(self):
        # Programın birden fazla açılmasını engelleme (Singleton Pattern)
        if not self.is_instance_the_only_one():
            messagebox.showerror(get_text('error'), get_text('app_already_running'))
            self.quit()
            return  # Programı kapat

        super().__init__()
        self.load_config()

        # Çoklu dil sistemi
        self.current_language = self.config.get('ui_language', 'tr')

        self.db = DatabaseManager()

        # Yeni modüller
        self.file_manager = FileManager(self.config)
        self.sound_manager = SoundManager()
        self.system_tray = SystemTray(self)

        self.setup_ui()

        # Kayıt durumu için değişken ekleyin
        self.recording_status = None
        self.recording_thread = None

        # Otomatik temizlik başlat
        self.file_manager.auto_cleanup_thread(callback=self.update_status)

        # Sistem trayı kur
        self.setup_system_tray()

        # Veritabanı durumunu kontrol et
        self.after(1000, self.check_database_status)

        # Güncelleme kontrolü
        self.after(2000, self.check_updates_startup)

        # self = FADIM (ana pencere)
        self.grid_rowconfigure(0, weight=1)  # main_container büyüsün
        self.grid_rowconfigure(1, weight=0)  # footer sabit
        self.grid_columnconfigure(0, weight=1)

        self.bind_hotkeys()

        # Standart dil haritası
        self.language_name_map = {
            "Türkçe": "tur",
            "İngilizce": "eng", 
            "Arapça": "ara",
            "Turkish": "tur",
            "English": "eng",
            "Arabic": "ara"
        }

        # Pencere minimize kontrolü
        self.protocol("WM_DELETE_WINDOW", self.on_window_close)

    def get_text(self, key, **kwargs):
        """Mevcut dilde metni al"""
        return get_text(key, self.current_language, **kwargs)

    def update_ui_language(self, event=None):
        """Arayüz dilini değiştir ve tüm UI'yi güncelle"""
        selected_lang = self.ui_lang_combo.get()
        lang_map = {'Türkçe': 'tr', 'English': 'en', 'العربية': 'ar'}
        new_lang = lang_map.get(selected_lang, 'tr')

        if new_lang != self.current_language:
            self.current_language = new_lang
            self.config['ui_language'] = new_lang

            # Config'i kaydet
            config_path = os.path.expanduser(
                os.path.join(self.config['base_directory'], self.config['directories']['config'], "config.json")
            )
            try:
                with open(config_path, 'w', encoding='utf-8') as f:
                    json.dump(self.config, f, indent=4, ensure_ascii=False)

                # UI'yi yeniden oluştur
                self.refresh_ui()
                self.status_var.set(self.get_text('language_selected', lang=selected_lang))

            except Exception as e:
                self.status_var.set(self.get_text('language_update_error', error=str(e)))

    def refresh_ui(self):
        """Tüm UI'yi yeniden oluştur (dil değişikliği için)"""
        # Ana pencereyi temizle ve yeniden oluştur
        for widget in self.winfo_children():
            widget.destroy()

        # UI'yi yeniden kur
        self.setup_ui()

        # Hotkey'leri yeniden bağla
        self.bind_hotkeys()

        # Sistem trayını yeniden kur
        if hasattr(self, 'system_tray'):
            self.system_tray.stop_tray()
            self.setup_system_tray()

    def get_mode_text(self, mode):
        """Mod için görüntü metni al"""
        mode_texts = {
            "ocr": self.get_text('mode_ocr'), 
            "gif": self.get_text('mode_gif'), 
            "screenshot": self.get_text('mode_screenshot')
        }
        return mode_texts.get(mode, self.get_text('mode_unknown'))

    def update_capture_mode(self):
        """Yakalama modunu güncelle"""
        mode = self.capture_mode.get()
        if mode == "ocr":
            self.config['auto_ocr'] = True
            self.config['gif_capture'] = False
        elif mode == "gif":
            self.config['auto_ocr'] = False
            self.config['gif_capture'] = True
        else:  # screenshot
            self.config['auto_ocr'] = False
            self.config['gif_capture'] = False

        

        # Config dosyasını kaydet
        config_path = os.path.expanduser(
            os.path.join(self.config['base_directory'], self.config['directories']['config'], "config.json")
        )
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)
            mode_text = self.get_mode_text(mode)
            self.status_var.set(self.get_text('mode_selected', mode=mode_text))
            # Mod göstergeyi güncelle
            if hasattr(self, 'mode_indicator'):
                self.mode_indicator.config(text=self.get_text('active_mode', mode=mode_text))
        except Exception as e:
            self.status_var.set(self.get_text('mode_change_error', error=str(e)))

    

    def quit_app(self):
        """Programı güvenli şekilde kapat"""
        if messagebox.askyesno(self.get_text('exit_title'), self.get_text('confirm_exit')):
            try:
                self.db.session.close()
            except:
                pass
            try:
                if hasattr(self, 'system_tray') and self.system_tray:
                    self.system_tray.stop_tray()
            except:
                pass
            self.destroy()
            import sys
            sys.exit(0)

    def load_config(self):
        # Belgeler/fadim/config klasörünü ayarla
        base_config_dir = os.path.expanduser('~/Documents/fadim/config')
        os.makedirs(base_config_dir, exist_ok=True)  # Yoksa oluştur

        # Config dosyasının tam yolu
        config_path = os.path.join(base_config_dir, 'config.json')

        # Varsayılan yapı
        default_config = {
            "language": "tur",
            "ui_language": "tr",
            "auto_ocr": True,
            "gif_capture": False,
            "click_sound": True,
            "show_cursor": False,
            "cursor_trail": False,
            "capture_hotkey": "ctrl+shift+f",
            "stop_hotkey": "esc",
            "base_directory": "~/Documents/fadim",
            "directories": {
                "screenshots": "screenshots",
                "documents": "documents",
                "logs": "logs",
                "config": "config"
            },
            "log_file": "logs/app.log",
            "cleanup": {
                "enabled": True,
                "days_to_keep": 30,
                "max_files": 100,
                "auto_cleanup_interval": 24
            }
        }

        # Dosya yoksa oluştur
        if not os.path.exists(config_path):
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(default_config, f, indent=4, ensure_ascii=False)

        # Dosyayı oku
        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = json.load(f)

        # Ana klasör ve alt klasörleri oluştur
        try:
            base_dir = os.path.expanduser(self.config['base_directory'])
            for dir_name in self.config['directories'].values():
                full_path = os.path.join(base_dir, dir_name)
                os.makedirs(full_path, exist_ok=True)

            # Log sistemi
            log_path = os.path.join(base_dir, self.config['directories']['logs'], "app.log")
            logging.basicConfig(
                filename=log_path,
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s',
                encoding='utf-8'
            )
            self.logger = logging.getLogger()
            self.logger.info("FADIM başlatıldı.")
        except Exception as e:
            print(f"❌ Klasör oluşturma hatası: {e}")
            # Varsayılan logger oluştur
            logging.basicConfig(level=logging.INFO)
            self.logger = logging.getLogger()

    def setup_ui(self):
        self.title("FADIM - Fast Access Digital Image-to-Text Manager")
        self.geometry("900x700")

        self.grid_rowconfigure(0, weight=1)
        self.grid_columnconfigure(0, weight=1)

        # Ana container
        main_container = ttk.Frame(self)
        main_container.grid(row=0, column=0, sticky='nsew')

        main_container.grid_rowconfigure(0, weight=0)  # header sabit
        main_container.grid_rowconfigure(1, weight=1)  # içerik esnek
        main_container.grid_rowconfigure(2, weight=0)  # footer sabit
        main_container.grid_columnconfigure(0, weight=1)

        # Header
        self.header_frame = self.create_header(main_container)
        self.header_frame.grid(row=0, column=0, sticky='ew')

        # Menü çubuğu
        menubar = tk.Menu(self)
        self['menu'] = menubar
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=self.get_text('menu_file'), menu=file_menu)
        file_menu.add_command(label=self.get_text('menu_open_image'), command=self.open_file)
        file_menu.add_command(label=self.get_text('menu_show_records'), command=self.show_records)

        # Ayarlar menüsü
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=self.get_text('menu_settings'), menu=settings_menu)
        settings_menu.add_command(label=self.get_text('menu_edit_settings'), command=self.open_settings_window)
        settings_menu.add_separator()
        settings_menu.add_command(label=self.get_text('menu_cleanup'), command=self.manual_cleanup)

        # Yardım menüsü
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label=self.get_text('menu_help'), menu=help_menu)
        help_menu.add_command(label=self.get_text('menu_usage_guide'), command=self.show_help)
        help_menu.add_command(label=self.get_text('menu_tesseract_install'), command=self.show_tesseract_help)
        help_menu.add_separator()
        help_menu.add_command(label=self.get_text('menu_update_check'), command=self.manual_update_check)
        help_menu.add_command(label=self.get_text('menu_website'), command=lambda: webbrowser.open(WEBSITE_URL))
        help_menu.add_separator()
        help_menu.add_command(label=self.get_text('menu_about'), command=self.show_about)

        # Ana içerik bölgesi
        main_frame = ttk.Frame(main_container)
        main_frame.grid(row=1, column=0, sticky='nsew', padx=10, pady=5)
        main_frame.grid_rowconfigure(1, weight=1)
        main_frame.grid_columnconfigure(0, weight=1)

        # Kontrol paneli
        control_frame = ttk.Frame(main_frame)
        control_frame.grid(row=0, column=0, sticky='ew')

        # UI Dil seçimi (solda)
        ttk.Label(control_frame, text=self.get_text('language_selection'), font=('Arial', 9, 'bold')).pack(side='left')
        self.ui_lang_combo = ttk.Combobox(control_frame, values=['Türkçe', 'English', 'العربية'], 
                                         state='readonly', width=10)
        ui_lang_names = {'tr': 'Türkçe', 'en': 'English', 'ar': 'العربية'}
        self.ui_lang_combo.set(ui_lang_names.get(self.current_language, 'Türkçe'))
        self.ui_lang_combo.pack(side='left', padx=5)
        self.ui_lang_combo.bind('<<ComboboxSelected>>', self.update_ui_language)

        ttk.Label(control_frame, text=self.get_text('ocr_language'), font=('Arial', 9, 'bold')).pack(side='left', padx=(15,5))
        # OCR dil seçimi - mevcut ayarı kullan
        current_lang_code = self.config.get('language', 'tur')
        lang_code_to_name = {'tur': self.get_text('lang_turkish'), 'eng': self.get_text('lang_english'), 'ara': self.get_text('lang_arabic')}

        self.lang_combo = ttk.Combobox(control_frame, values=[self.get_text('lang_turkish'), self.get_text('lang_english'), self.get_text('lang_arabic')], 
                                      state='readonly', width=12)
        self.lang_combo.set(lang_code_to_name.get(current_lang_code, self.get_text('lang_turkish')))
        self.lang_combo.pack(side='left', padx=5)
        self.lang_combo.bind('<<ComboboxSelected>>', self.update_language_setting)
        self.CreateToolTip(self.lang_combo, self.get_text('tooltip_lang_combo'))
        ttk.Label(control_frame, text=self.get_text('help_text'), font=('Arial', 9, 'bold')).pack(side='left', padx=(15,5))

        # Dinamik kısayol metni
        current_hotkey = self.config.get('capture_hotkey', 'ctrl+shift+f').upper()
        help_text = self.get_text('help_hotkey', hotkey=current_hotkey)
        self.help_label = ttk.Label(control_frame, text=help_text)
        self.help_label.pack(side='left')

        # OCR Sonuç alanı
        result_frame = ttk.LabelFrame(main_frame, text=self.get_text('ocr_result'))
        result_frame.grid(row=1, column=0, sticky='nsew', padx=10, pady=5)
        result_frame.grid_rowconfigure(1, weight=1)
        result_frame.grid_columnconfigure(0, weight=1)

        # Araç çubuğu
        toolbar_frame = ttk.Frame(result_frame)
        toolbar_frame.grid(row=0, column=0, sticky='ew', padx=5, pady=5)

        # Font boyutu kontrolü
        font_frame = ttk.Frame(toolbar_frame)
        font_frame.pack(side='left', padx=5)
        ttk.Label(font_frame, text=self.get_text('font_size'), font=('Arial', 9, 'bold')).pack(side='top')
        self.font_size = ttk.Spinbox(font_frame, from_=8, to=40, width=4, command=self.change_font_size)
        self.font_size.set(12)
        self.font_size.pack(side='top')
        self.CreateToolTip(self.font_size, self.get_text('tooltip_font_size'))

        # Ayırıcı
        ttk.Separator(toolbar_frame, orient='vertical').pack(side='left', fill='y', padx=10)

        # Hizalama butonları
        align_frame = ttk.Frame(toolbar_frame)
        align_frame.pack(side='left', padx=5)
        ttk.Label(align_frame, text=self.get_text('alignment'), font=('Arial', 9, 'bold')).pack(side='top')
        align_buttons_frame = ttk.Frame(align_frame)
        align_buttons_frame.pack(side='top')

        align_data = [
            (self.get_text('align_left'), 'left', self.get_text('tooltip_align_left')),
            (self.get_text('align_center'), 'center', self.get_text('tooltip_align_center')),
            (self.get_text('align_right'), 'right', self.get_text('tooltip_align_right'))
        ]

        for text, align, tooltip in align_data:
            btn = ttk.Button(align_buttons_frame, text=text, width=8, 
                           command=lambda a=align: self.align_text(a))
            btn.pack(side='left', padx=1)
            self.CreateToolTip(btn, tooltip)

        # Ayırıcı
        ttk.Separator(toolbar_frame, orient='vertical').pack(side='left', fill='y', padx=10)

        # Stil butonları
        style_frame = ttk.Frame(toolbar_frame)
        style_frame.pack(side='left', padx=5)
        ttk.Label(style_frame, text=self.get_text('style'), font=('Arial', 9, 'bold')).pack(side='top')
        style_buttons_frame = ttk.Frame(style_frame)
        style_buttons_frame.pack(side='top')

        self.bold_var = tk.BooleanVar()
        self.italic_var = tk.BooleanVar()

        bold_btn = ttk.Checkbutton(style_buttons_frame, text=self.get_text('bold'), width=10, 
                                  variable=self.bold_var, command=self.apply_style)
        bold_btn.pack(side='left', padx=2)
        self.CreateToolTip(bold_btn, self.get_text('tooltip_bold'))

        italic_btn = ttk.Checkbutton(style_buttons_frame, text=self.get_text('italic'), width=10, 
                                    variable=self.italic_var, command=self.apply_style)
        italic_btn.pack(side='left', padx=2)
        self.CreateToolTip(italic_btn, self.get_text('tooltip_italic'))

        # Text Container
        text_container = tk.Frame(result_frame, height=300)
        text_container.grid(row=1, column=0, sticky='nsew', padx=5, pady=5)
        text_container.grid_propagate(False)
        self.result_text = scrolledtext.ScrolledText(text_container, wrap=tk.WORD, font=("Arial", 12))
        self.result_text.pack(fill='both', expand=True)
        self.result_text.bind('<KeyRelease>', self.update_stats)
        self.result_text.bind('<Button-1>', self.update_stats)

        # İstatistik alanı
        self.create_stats_area(result_frame)

        # Yakalama Modu ve GIF Seçenekleri
        options_frame = ttk.Frame(main_frame)
        options_frame.grid(row=2, column=0, sticky='ew', pady=5)

        # Sol taraf - Yakalama Modu
        mode_frame = ttk.LabelFrame(options_frame, text=self.get_text('capture_mode'))
        mode_frame.pack(side='left', fill='y', padx=(0, 5))

        # Mod belirleme
        current_mode = "ocr"
        if self.config.get('gif_capture', False):
            current_mode = "gif"
        elif not self.config.get('auto_ocr', True) and not self.config.get('gif_capture', False):
            current_mode = "screenshot"

        self.capture_mode = tk.StringVar(value=current_mode)

        # Mod seçenekleri yatay olarak sıralanmış
        mode_options_frame = ttk.Frame(mode_frame)
        mode_options_frame.pack(padx=10, pady=5)

        ocr_radio = ttk.Radiobutton(mode_options_frame, text=self.get_text('ocr_mode'), 
                                   variable=self.capture_mode, value="ocr",
                                   command=self.update_capture_mode)
        ocr_radio.pack(side='left', padx=(0, 10))
        self.CreateToolTip(ocr_radio, self.get_text('tooltip_ocr_mode'))

        gif_radio = ttk.Radiobutton(mode_options_frame, text=self.get_text('gif_mode'), 
                                   variable=self.capture_mode, value="gif",
                                   command=self.update_capture_mode)
        gif_radio.pack(side='left', padx=(0, 10))
        self.CreateToolTip(gif_radio, self.get_text('tooltip_gif_mode'))

        screenshot_radio = ttk.Radiobutton(mode_options_frame, text=self.get_text('screenshot_mode'), 
                                          variable=self.capture_mode, value="screenshot",
                                          command=self.update_capture_mode)
        screenshot_radio.pack(side='left')
        self.CreateToolTip(screenshot_radio, self.get_text('tooltip_screenshot_mode'))

        # Aktif mod göstergesi
        self.mode_indicator = ttk.Label(mode_frame, text=self.get_text('active_mode', mode=self.get_mode_text(current_mode)), 
                                      font=('Arial', 9, 'bold'), foreground='green')
        self.mode_indicator.pack(padx=10, pady=(5,10))

        # Sağ taraf - GIF Seçenekleri
        gif_options_frame = ttk.LabelFrame(options_frame, text=self.get_text('gif_options'))
        gif_options_frame.pack(side='left', fill='both', expand=True, padx=(5, 0))

        # GIF seçenekleri container
        gif_container = ttk.Frame(gif_options_frame)
        gif_container.pack(fill='x', padx=10, pady=5)

        # GIF kalite seçimi
        quality_frame = ttk.Frame(gif_container)
        quality_frame.pack(side='left', padx=(0, 10))
        ttk.Label(quality_frame, text=self.get_text('quality')).pack(side='left')
        self.gif_quality = ttk.Combobox(quality_frame, values=['low', 'medium', 'high'], 
                                       state='readonly', width=8)
        self.gif_quality.set('medium')
        self.gif_quality.pack(side='left', padx=5)
        self.CreateToolTip(self.gif_quality, self.get_text('tooltip_gif_quality'))

        # GIF süre seçimi
        duration_frame = ttk.Frame(gif_container)
        duration_frame.pack(side='left', padx=(0, 10))
        ttk.Label(duration_frame, text=self.get_text('duration')).pack(side='left')
        self.gif_duration = ttk.Combobox(duration_frame, values=[5, 10, 15, 20, 30, 45, 60], 
                                          state='readonly', width=4)
        self.gif_duration.set(10)
        self.gif_duration.pack(side='left', padx=5)
        self.CreateToolTip(self.gif_duration, self.get_text('tooltip_gif_duration'))

        # GIF FPS seçimi
        fps_frame = ttk.Frame(gif_container)
        fps_frame.pack(side='left', padx=(0, 10))
        ttk.Label(fps_frame, text=self.get_text('fps')).pack(side='left')
        self.gif_fps = ttk.Combobox(fps_frame, values=[5, 10, 15, 20, 25, 30], 
                                     state='readonly', width=4)
        self.gif_fps.set(10)
        self.gif_fps.pack(side='left', padx=5)
        self.CreateToolTip(self.gif_fps, self.get_text('tooltip_gif_fps'))

        # Bölge seçimi butonu
        region_btn = ttk.Button(gif_container, text=self.get_text('select_region'), 
                               command=self.show_region_selector)
        region_btn.pack(side='left', padx=(0, 10))
        self.CreateToolTip(region_btn, self.get_text('tooltip_select_region'))

        # Butonlar
        button_frame = ttk.LabelFrame(main_frame, text=self.get_text('operations'))
        button_frame.grid(row=3, column=0, sticky='ew', pady=5)

        # Sol taraftaki butonlar
        # Ana butonlar
        btn_take_screenshot = ttk.Button(button_frame, text=self.get_text('take_screenshot'), command=self.capture_and_process)
        btn_take_screenshot.pack(side='left', padx=5, pady=5)
        self.CreateToolTip(btn_take_screenshot, self.get_text('tooltip_take_screenshot'))

        btn_open_file = ttk.Button(button_frame, text=self.get_text('open_file'), command=self.open_file)
        btn_open_file.pack(side='left', padx=5, pady=5)
        self.CreateToolTip(btn_open_file, self.get_text('tooltip_open_file'))

        # GIF dur butonu - başlangıçta pasif
        self.btn_stop_gif = ttk.Button(button_frame, text=self.get_text('stop_gif'), command=self.stop_gif_recording, state='disabled')
        self.btn_stop_gif.pack(side='left', padx=5, pady=5)
        self.CreateToolTip(self.btn_stop_gif, self.get_text('tooltip_stop_gif'))

        # Diğer butonlar
        other_buttons_data = [
            (self.get_text('save_txt'), lambda: self.save_text('txt'), self.get_text('tooltip_save_txt')),
            (self.get_text('save_docx'), lambda: self.save_text('docx'), self.get_text('tooltip_save_docx')),
            (self.get_text('clear_text'), lambda: self.result_text.delete(1.0, tk.END), self.get_text('tooltip_clear_text')),
            (self.get_text('copy_text'), self.copy_text, self.get_text('tooltip_copy_text'))
        ]

        for label, cmd, tooltip in other_buttons_data:
            btn = ttk.Button(button_frame, text=label, command=cmd)
            btn.pack(side='left', padx=5, pady=5)
            self.CreateToolTip(btn, tooltip)

        # Kapat butonunu sağa yasla
        close_btn = ttk.Button(button_frame, text=self.get_text('close_app'), command=self.quit_app)
        close_btn.pack(side='right', padx=5, pady=5)
        self.CreateToolTip(close_btn, self.get_text('tooltip_close_app'))

        # Durum çubuğu
        status_frame = ttk.Frame(main_frame)
        status_frame.grid(row=4, column=0, sticky='ew', pady=5)
        self.status_var = tk.StringVar(value=self.get_text('status_ready'))
        ttk.Label(status_frame, textvariable=self.status_var).pack(side='left', pady=5)
        self.open_docs_button = ttk.Button(status_frame, text="📁", width=3, command=self.open_save_directory)
        self.open_docs_button.pack(side='left', padx=5)
        self.CreateToolTip(self.open_docs_button, self.get_text('tooltip_docs_folder'))

        self.open_gif_button = ttk.Button(status_frame, text="🎞️", width=3, command=self.open_gif_directory)
        self.open_gif_button.pack(side='left', padx=5)
        self.CreateToolTip(self.open_gif_button, self.get_text('tooltip_gif_folder'))

        # Footer
        self.footer_frame = self.create_footer(main_container)
        self.footer_frame.grid(row=2, column=0, sticky='ew')

    def update_language_setting(self, event=None):
        selected_lang = self.lang_combo.get()
        # Tesseract için dil koduna çevir
        tesseract_lang = self.language_name_map.get(selected_lang, 'eng')
        self.config['language'] = tesseract_lang

        # Config dosyasına kalıcı olarak yaz
        config_path = os.path.expanduser(
            os.path.join(self.config['base_directory'], self.config['directories']['config'], "config.json")
        )
        try:
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)

            # Dil özel uyarıları
            if selected_lang == 'Arapça':
                self.status_var.set(f"🌐 {selected_lang} seçildi - Arapça metinler sağa yaslanacak")
            elif selected_lang == 'Türkçe':
                self.status_var.set(f"🌐 {selected_lang} seçildi - Türkçe karakterler tanınacak (ç,ğ,ı,ö,ş,ü)")
            else:
                self.status_var.set(f"🌐 {selected_lang} seçildi")

        except Exception as e:
            self.status_var.set(f"❌ Dil güncellenemedi: {str(e)}")

    def get_resource_path(self, relative_path):
        """PyInstaller derlemesi için uygun dosya yolunu döndürür."""
        try:
            base_path = sys._MEIPASS
        except AttributeError:
            base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)

    def create_header(self, parent):
        frame = ttk.Frame(parent, relief='solid', borderwidth=1)
        frame.grid(row=0, column=0, sticky='ew', padx=5, pady=5)

        icon_title_frame = ttk.Frame(frame)
        icon_title_frame.pack(side='left', padx=10, pady=10)

        try:
            icon_path = self.get_resource_path(os.path.join('core', 'fadim_icon.png'))
            print(f"[DEBUG] İkon yolu: {icon_path}")

            if os.path.exists(icon_path):
                icon_image = Image.open(icon_path).resize((42, 42), Image.Resampling.LANCZOS)
                self.icon_photo = ImageTk.PhotoImage(icon_image)
                ttk.Label(icon_title_frame, image=self.icon_photo).grid(row=0, column=0, rowspan=2, sticky='n')
            else:
                print("[DEBUG] İkon bulunamadı.")
        except Exception as e:
            print(f"[HATA] İkon yüklenemedi: {e}")

        # Metin bloğu: ikonla aynı hizada, sola yaslı
        title_text_frame = ttk.Frame(icon_title_frame)
        title_text_frame.grid(row=0, column=1, sticky='w', padx=(10, 0))

        # Dil ayarına göre başlık göster
        if self.current_language == 'tr':
            # Türkçe modunda: İngilizce + Türkçe başlık
            en_line = ttk.Frame(title_text_frame)
            en_line.pack(anchor='w')
            ttk.Label(en_line, text="FADIM — Fast Access Digital Image-to-Text Manager", font=('Arial', 13)).pack(side='left')
            ttk.Label(en_line, text="  Capture, Read, Save — Instantly!", font=('Arial', 10, 'italic')).pack(side='left')

            tr_line = ttk.Frame(title_text_frame)
            tr_line.pack(anchor='w')
            ttk.Label(tr_line, text="Fadim — Fotoğraftan Anında Dijital Metin", font=('Arial', 13)).pack(side='left')
            ttk.Label(tr_line, text="  Yakala, Oku, Kaydet — Anında!", font=('Arial', 10, 'italic')).pack(side='left')

        elif self.current_language == 'ar':
            # Arapça modunda: sadece Arapça başlık
            ar_line = ttk.Frame(title_text_frame)
            ar_line.pack(anchor='w')
            ttk.Label(ar_line, text="فاديم — مدير النص الرقمي للوصول السريع للصور", font=('Arial', 13)).pack(side='left')
            ttk.Label(ar_line, text="  التقط، اقرأ، احفظ — فوراً!", font=('Arial', 10, 'italic')).pack(side='left')

        else:  # İngilizce modu
            # İngilizce modunda: sadece İngilizce başlık
            en_line = ttk.Frame(title_text_frame)
            en_line.pack(anchor='w')
            ttk.Label(en_line, text="FADIM — Fast Access Digital Image-to-Text Manager", font=('Arial', 13)).pack(side='left')
            ttk.Label(en_line, text="  Capture, Read, Save — Instantly!", font=('Arial', 10, 'italic')).pack(side='left')

        return frame

    def create_stats_area(self, parent):
        """İstatistik alanı oluşturma"""
        stats_frame = ttk.Frame(parent)
        stats_frame.grid(row=2, column=0, sticky='ew', padx=5, pady=2)

        self.stats_var = tk.StringVar()
        # Başlangıç değeri için dil sistemi kullan
        initial_stats = f"{self.get_text('stats_lines')}: 0 | {self.get_text('stats_words')}: 0 | {self.get_text('stats_characters')}: 0"
        self.stats_var.set(initial_stats)

        stats_label = ttk.Label(stats_frame, textvariable=self.stats_var, font=('Arial', 9))
        stats_label.pack(side='left')

    def create_footer(self, parent):
        """Footer bölümü oluşturma"""
        footer_frame = ttk.Frame(parent, relief='solid', borderwidth=1)
        footer_content = ttk.Frame(footer_frame)
        footer_content.pack(pady=5)

        # Program isimleri - dil bazında göster
        if self.current_language == 'tr':
            program_text = "FADIM - Fast Access Digital Image-to-Text Manager | Fotoğraftan Anında Dijital Metin"
        elif self.current_language == 'en':
            program_text = "FADIM - Fast Access Digital Image-to-Text Manager"
        else:  # Arabic
            program_text = "فاديم — مدير النص الرقمي للوصول السريع للصور"

        program_names = ttk.Label(footer_content, text=program_text, font=('Arial', 10))
        program_names.pack()

        # Versiyon ve linkler çerçevesi
        links_frame = ttk.Frame(footer_content)
        links_frame.pack()

        # Versiyon bilgisi
        version_label = ttk.Label(links_frame, text=f"v{CURRENT_VERSION}", 
                                font=('Arial', 9), foreground='gray')
        version_label.pack(side='left', padx=(0, 10))

        # Website linki
        website_label = ttk.Label(links_frame, text="🌐", 
                                font=('Arial', 12), 
                                foreground='blue', cursor='hand2')
        website_label.pack(side='left', padx=(0, 10))
        website_label.bind("<Button-1>", lambda e: webbrowser.open(WEBSITE_URL))
        self.CreateToolTip(website_label, self.get_text('tooltip_website'))

        # Ana site linki
        main_site_label = ttk.Label(links_frame, text="🏠", 
                                font=('Arial', 12), 
                                foreground='blue', cursor='hand2')
        main_site_label.pack(side='left')
        main_site_label.bind("<Button-1>", lambda e: webbrowser.open("https://www.muallimun.net"))
        self.CreateToolTip(main_site_label, self.get_text('tooltip_main_site'))

        return footer_frame

    def update_stats(self, event=None):
        """Metin istatistiklerini güncelleme"""
        text = self.result_text.get(1.0, tk.END)

        # Satır sayısı
        lines = len(text.split('\n')) - 1  # Son boş satırı çıkar

        # Sözcük sayısı
        words = len(text.split())

        # Karakter sayısı (boşluklar dahil, son newline hariç)
        chars = len(text.rstrip('\n'))

        # Dil sistemini kullanarak formatla
        stats_text = f"{self.get_text('stats_lines')}: {lines} | {self.get_text('stats_words')}: {words} | {self.get_text('stats_characters')}: {chars}"
        self.stats_var.set(stats_text)

    def change_font_size(self, event=None):
        """Font boyutunu anında değiştirme"""
        try:
            size = int(self.font_size.get())
            current_font = self.result_text['font']
            if isinstance(current_font, str):
                font_parts = current_font.split()
                font_family = font_parts[0] if font_parts else 'TkDefaultFont'
            else:
                font_family = 'TkDefaultFont'

            # Mevcut stil özelliklerini koru
            weight = 'bold' if self.bold_var.get() else 'normal'
            slant = 'italic' if self.italic_var.get() else 'roman'

            self.result_text.configure(font=(font_family, size, weight, slant))
        except ValueError:
            pass  # Geçersiz değer girilirse sessizce geç

    def align_text(self, align):
        self.result_text.tag_configure('align', justify=align)
        self.result_text.tag_add('align', '1.0', 'end')

    def apply_style(self):
        current_font = self.result_text['font']
        if isinstance(current_font, str):
            font_parts = current_font.split()
            size = font_parts[1] if len(font_parts) > 1 else self.font_size.get()
            font_family = font_parts[0] if font_parts else 'TkDefaultFont'
        else:
            size = self.font_size.get()
            font_family = 'TkDefaultFont'

        weight = 'bold' if self.bold_var.get() else 'normal'
        slant = 'italic' if self.italic_var.get() else 'roman'

        self.result_text.configure(font=(font_family, size, weight, slant))

    def open_save_directory(self):
        import subprocess, platform

        # Belgeler/fadim/documents klasörünü hedefle
        documents_dir = os.path.join(
            os.path.expanduser(self.config['base_directory']),
            self.config['directories']['documents']
        )

        if platform.system() == "Windows":
            os.startfile(documents_dir)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", documents_dir])
        else:
            subprocess.Popen(["xdg-open", documents_dir])

    def open_gif_directory(self):
        import subprocess, platform

        gif_dir = os.path.join(
            os.path.expanduser(self.config['base_directory']),
            self.config['directories']['screenshots']
        )

        if platform.system() == "Windows":
            os.startfile(gif_dir)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", gif_dir])
        else:
            subprocess.Popen(["xdg-open", gif_dir])

    def bind_hotkeys(self):
        """Hotkey bağlama - Kullanıcı tanımlı kısayollarla"""
        if keyboard is None:
            self.status_var.set(self.get_text('keyboard_not_found'))
            return
            
        try:
            # Önceki hotkey'leri temizle
            try:
                keyboard.unhook_all_hotkeys()
            except:
                pass

            # Yakalama hotkey'i
            capture_hotkey = self.config.get('capture_hotkey', 'ctrl+shift+f')

            def hotkey_capture():
                """Thread-safe capture işlemi"""
                self.after(0, self.capture_and_process)

            keyboard.add_hotkey(capture_hotkey, hotkey_capture)

            self.status_var.set(self.get_text('hotkeys_active', hotkey=capture_hotkey))
            print(f"✅ Hotkey bağlandı: {capture_hotkey}")

        except ImportError:
            self.status_var.set(self.get_text('keyboard_not_found'))
            print("❌ keyboard kütüphanesi bulunamadı")
        except Exception as e:
            self.status_var.set(self.get_text('hotkey_error', error=str(e)))
            print(f"❌ Hotkey hatası: {e}")
            # Hata durumunda sadece GUI ile devam et

    def is_instance_the_only_one(self):
        """Sadece bir örneğin çalıştığından emin ol"""
        self.app_id = "fadim_app_id"  # Benzersiz bir ID oluştur
        self.instance = platform.node()
        try:
            if platform.system() == "Windows":
                import win32event, win32api
                self.mutex = win32event.CreateMutex(None, False, self.app_id)
                self.last_error = win32api.GetLastError()
                if self.last_error == 183:  # Zaten varsa
                    return False
        except ImportError:
            print("⚠️ Windows-specific libraries not available")
        except Exception as e:
            print(f"⚠️ Instance check error: {e}")
        return True

    def capture_and_process(self):
        # GIF kaydı ayarı varsa, sadece kayıt yap
        if self.config.get('gif_capture', False):
            save_dir = os.path.join(os.path.expanduser(self.config['base_directory']), self.config['directories']['screenshots'])
            os.makedirs(save_dir, exist_ok=True)

            timestamp = time.strftime('%Y%m%d_%H%M%S')
            output_gif = os.path.join(save_dir, f"capture_{timestamp}.gif")

            # Seçilen bölgeyi kontrol et
            region = getattr(self, 'selected_region', None)

            self.status_var.set(self.get_text('gif_recording_starting'))

            # Ekranın sol üst köşesinde uyarı penceresi oluştur
            self.recording_overlay = tk.Toplevel()
            self.recording_overlay.title("Kayıt Durumu")
            self.recording_overlay.geometry("220x60+10+10")  # +10+10 = sol üst köşe
            self.recording_overlay.resizable(False, False)
            self.recording_overlay.attributes('-topmost', True)  # Her zaman üstte
            self.recording_overlay.configure(bg='yellow')

            # Pencere kenarlığını kaldır ve minimize/close butonlarını gizle
            self.recording_overlay.overrideredirect(True)

            label = tk.Label(self.recording_overlay, 
                           text=self.get_text('recording_status'),
                           bg="yellow", fg="black", font=("Arial", 9, "bold"),
                           justify='center')
            label.pack(expand=True, fill='both', padx=5, pady=5)

            # Kayıt durumu değişkenini oluştur
            from core.gif_capture import RecordingStatus
            self.recording_status = RecordingStatus()

            # Global ESC tuşu için thread
            def monitor_escape():
                try:
                    import keyboard
                    while self.recording_status and not self.recording_status.should_stop():
                        if keyboard.is_pressed('esc'):
                            self.stop_recording()
                            # Overlay'i kapat
                            if hasattr(self, 'recording_overlay'):
                                try:
                                    if self.recording_overlay.winfo_exists():
                                        self.recording_overlay.after(0, self.recording_overlay.destroy)
                                except:
                                    pass
                            break
                        time.sleep(0.1)
                except Exception as e:
                    print(f"ESC monitör hatası: {e}")

            # ESC monitör thread'ini başlat
            escape_thread = threading.Thread(target=monitor_escape, daemon=True)
            escape_thread.start()

            self.update()

            # Ses bildirimi
            self.sound_manager.play_sound('recording_start')

            # Thread'i ve durumunu sakla  
            self.recording_thread = threading.Thread(
                target=lambda: self.run_recording(output_gif, region),
                daemon=True
            )
            self.recording_thread.start()

            # Kayıt durumunu kontrol et
            self.check_recording_status(output_gif)

            # Bölge seçimini reset et
            if hasattr(self, 'selected_region'):
                self.selected_region = None

            return  # OCR yapılmasın, sadece GIF

        # Ekran alıntısı alınıyor
        self.status_var.set(self.get_text('capture_starting'))
        self.update()
        self.withdraw()

        def capture_in_thread():
            try:
                screenshot_path = capture_screen_region(
                    play_sound=self.config.get('click_sound', True),
                    show_cursor=self.config.get('show_cursor', False)
                )
                if screenshot_path:
                    if self.config.get("auto_ocr", True):
                        self.after(100, lambda: self.process_capture_result(screenshot_path))
                    else:
                        self.after(100, lambda: self.show_window())
                        self.status_var.set("Ekran görüntüsü alındı (OCR yapılmadı)")
                else:
                    self.after(100, self.show_window)
            except Exception as e:
                self.after(100, lambda: self.handle_capture_error(str(e)))

        self.after(200, lambda: threading.Thread(target=capture_in_thread, daemon=True).start())

    def run_recording(self, output_path, region=None):
        """GIF kaydını gerçekleştiren metod"""
        from core.gif_capture import record_gif

        # GIF dur butonunu aktif et
        self.after(0, lambda: self.btn_stop_gif.config(state='normal'))

        # Kalite, süre ve FPS ayarlarını al
        quality = self.gif_quality.get()
        duration = int(self.gif_duration.get())
        fps = int(self.gif_fps.get())

        success = record_gif(
            output_path, 
            duration=duration,
            fps=fps, 
            region=region,
            status=self.recording_status,
            show_cursor=self.config.get('show_cursor', False),
            cursor_trail=self.config.get('cursor_trail', False),
            quality=quality
        )

        # Ses bildirimi
        self.sound_manager.play_sound('recording_stop' if success else 'error')

        # GIF dur butonunu pasif et
        self.after(0, lambda: self.btn_stop_gif.config(state='disabled'))

        # Kayıt tamamlandı, gerekli temizlikleri yap
        if hasattr(self, 'recording_overlay'):
            try:
                if self.recording_overlay.winfo_exists():
                    self.after(0, lambda: self.recording_overlay.destroy())
                    self.after(0, lambda: delattr(self, 'recording_overlay') if hasattr(self, 'recording_overlay') else None)
            except:
                pass
        self.recording_status = None

    def stop_recording(self):
        """Kullanıcı ESC tuşuna bastığında çağrılan metod"""
        if self.recording_status:
            self.recording_status.request_stop()
            self.status_var.set(self.get_text('gif_stop_requested'))
            # GIF dur butonunu pasif et
            if hasattr(self, 'btn_stop_gif'):
                self.btn_stop_gif.config(state='disabled')
        # Overlay penceresini güvenli şekilde kapat
        if hasattr(self, 'recording_overlay'):
            try:
                if self.recording_overlay.winfo_exists():
                    self.recording_overlay.destroy()
                    delattr(self, 'recording_overlay')
            except:
                pass

    def stop_gif_recording(self):
        """Buton ile GIF kaydını durdurma"""
        if self.recording_status:
            self.recording_status.request_stop()
            self.status_var.set(self.get_text('gif_stopped_button'))
            # GIF dur butonunu pasif et
            self.btn_stop_gif.config(state='disabled')
        else:
            self.status_var.set(self.get_text('no_active_gif'))

    def check_recording_status(self, output_path):
        """Kayıt durumunu periyodik olarak kontrol et"""
        if hasattr(self, 'recording_thread') and self.recording_thread.is_alive():
            # Kayıt hala devam ediyor, 500ms sonra tekrar kontrol et
            self.after(500, lambda: self.check_recording_status(output_path))
        else:
            # Kayıt bitti
            self.status_var.set(self.get_text('gif_saved', filename=os.path.basename(output_path)))
            # Overlay'i güvenli şekilde kapat
            if hasattr(self, 'recording_overlay'):
                try:
                    if self.recording_overlay.winfo_exists():
                        self.recording_overlay.destroy()
                        delattr(self, 'recording_overlay')
                except:
                    pass
            self.recording_status = None

    def process_capture_result(self, screenshot_path):
        """Capture sonucunu işleme"""
        self.show_window()
        self.process_image(screenshot_path)

    def handle_capture_error(self, error_msg):
        """Capture hatasını ele alma"""
        self.show_window()
        self.status_var.set(self.get_text('capture_error', error=error_msg))

    def show_window(self):
        """Pencereyi tekrar göster"""
        self.deiconify()
        self.lift()
        self.focus_force()

    def is_arabic_text(self, text):
        """Metnin gerçekten Arapça olup olmadığını belirler"""
        if not text:
            return False

        # En az 3 harfli Arapça kelimeleri bul
        arabic_words = re.findall(r'[\u0600-\u06FF]{3,}', text)

        # En az 3 Arapça kelime varsa Arapça kabul edilir
        return len(arabic_words) >= 3

    def process_image(self, image_path):
        try:
            # OCR işlemini gerçekleştir
            current_lang = self.lang_combo.get()

            # OCR fonksiyonundan tuple döner: (text, image)
            ocr_result = ocr_image(image_path, lang=current_lang)
            if isinstance(ocr_result, tuple) and len(ocr_result) == 2:
                text, image = ocr_result
            else:
                # Eski format uyumluluğu için
                text = ocr_result
                image = None

            # Boş text kontrolü - veritabanına kaydetmeden önce kontrol et
            if not text or text.strip() == "" or text.strip() == "OCR işlemi sonuç vermedi":
                text = "OCR işlemi sonuç vermedi"
                # Boş sonuç olsa bile UI'yi güncelle
                self.result_text.delete(1.0, tk.END)
                self.result_text.insert(tk.END, text)
                self.status_var.set(self.get_text('ocr_failed_general'))
                return

            # Başarılı OCR sonucu - UI'yi güncelle
            self.result_text.delete(1.0, tk.END)
            self.result_text.insert(tk.END, text)

            # Panoya kopyala (sadece başarılı durumda)
            if image is not None:
                if platform.system() == 'Windows':
                    import io
                    import win32clipboard
                    from PIL import ImageGrab

                    def send_to_clipboard(clip_type, data):
                        win32clipboard.OpenClipboard()
                        win32clipboard.EmptyClipboard()
                        win32clipboard.SetClipboardData(clip_type, data)
                        win32clipboard.CloseClipboard()

                    output = io.BytesIO()
                    image.convert("RGB").save(output, "BMP")
                    data = output.getvalue()[14:]
                    send_to_clipboard(win32clipboard.CF_DIB, data)
                    print("✅ Ekran görüntüsü panoya kopyalandı (Windows)")
                elif platform.system() == 'Darwin':  # macOS
                    try:
                        import subprocess
                        image.save('temp_screenshot.png')
                        subprocess.run(['osascript', '-e',
                                        'set the clipboard to POSIX file "temp_screenshot.png"'])
                        os.remove('temp_screenshot.png')  # Geçici dosyayı sil
                        print("✅ Ekran görüntüsü panoya kopyalandı (macOS)")
                    except Exception as e:
                        print(f"❌ macOS panoya kopyalama hatası: {e}")
                else:  # Linux
                    try:
                        import subprocess
                        image.save('temp_screenshot.png')
                        subprocess.run(['xclip', '-selection', 'clipboard', '-t', 'image/png', '-i',
                                        'temp_screenshot.png'])
                        os.remove('temp_screenshot.png')  # Geçici dosyayı sil
                        print("✅ Ekran görüntüsü panoya kopyalandı (Linux)")
                    except Exception as e:
                        print(f"❌ Linux panoya kopyalama hatası: {e}")

            # Arapça metin kontrolü ve sağa yaslama
            if self.is_arabic_text(text):
                self.align_text('right')
                self.status_var.set(self.get_text('arabic_detected'))
            else:
                self.align_text('left')

            self.update_stats()

            # Veritabanına kaydetme işlemi - sadece başarılı OCR sonuçları için
            try:
                # Dil kodunu Tesseract formatına çevir
                lang_map = {
                    "Türkçe": "tur",
                    "İngilizce": "eng", 
                    "Arapça": "ara",
                    "Turkish": "tur",
                    "English": "eng",
                    "Arabic": "ara"
                }
                tesseract_lang = lang_map.get(current_lang, 'tur')
                
                # OCR sonucunu veritabanına kaydet
                record_id = self.db.save_text(text, source='screenshot', language=tesseract_lang)
                self.logger.info(f"OCR sonucu veritabanına kaydedildi. ID: {record_id}, Dil: {tesseract_lang}")
                print(f"✅ Veritabanı kaydı başarılı: ID={record_id}, Text uzunluğu={len(text)}")

                # Durum çubuğunu güncelle
                if self.get_text('arabic_detected') not in self.status_var.get():
                    image_status = self.get_text('image_clipboard_success') if image is not None else self.get_text('image_clipboard_failed')
                    self.status_var.set(self.get_text('text_processed_saved', id=record_id, status=image_status))

            except Exception as db_error:
                self.logger.error(f"Veritabanı kayıt hatası: {str(db_error)}")
                print(f"❌ Veritabanı kayıt hatası: {str(db_error)}")
                self.status_var.set(self.get_text('text_processed_error', error=str(db_error)))

        except RuntimeError as e:
            messagebox.showerror(self.get_text('ocr_error'), str(e))
            self.status_var.set(self.get_text('ocr_failed_tesseract'))

        except Exception as e:
            messagebox.showerror(self.get_text('ocr_error'), f"{self.get_text('unexpected_error')}: {str(e)}")
            self.status_var.set(self.get_text('ocr_failed_general'))

    def open_file(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Image/PDF files", "*.png *.jpg *.jpeg *.pdf")])
        if file_path:
            self.process_image(file_path)

    def open_settings_window(self):
        win = tk.Toplevel(self)
        win.title(self.get_text('settings_title'))
        win.geometry("500x400")
        win.resizable(False, False)

        notebook = ttk.Notebook(win)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)

        # Genel sekmesi
        general_frame = ttk.Frame(notebook)
        notebook.add(general_frame, text=self.get_text('settings_general'))

        # UI Dil seçimi
        ui_lang_frame = ttk.LabelFrame(general_frame, text=self.get_text('language_selection'))
        ui_lang_frame.pack(fill='x', pady=5)

        ui_lang_names = {'tr': 'Türkçe', 'en': 'English', 'ar': 'العربية'}
        current_ui_lang = ui_lang_names.get(self.current_language, 'Türkçe')
        
        ui_lang_var = tk.StringVar(value=current_ui_lang)
        ui_lang_combo_settings = ttk.Combobox(ui_lang_frame, textvariable=ui_lang_var,
                                             values=['Türkçe', 'English', 'العربية'], state='readonly')
        ui_lang_combo_settings.pack(padx=10, pady=5)

        lang_frame = ttk.LabelFrame(general_frame, text=self.get_text('default_ocr_lang'))
        lang_frame.pack(fill='x', pady=5)

        # Mevcut ayarı dil adına çevir
        current_lang_code = self.config.get('language', 'tur')
        lang_code_to_name = {'tur': self.get_text('lang_turkish'), 'eng': self.get_text('lang_english'), 'ara': self.get_text('lang_arabic')}
        current_lang_name = lang_code_to_name.get(current_lang_code, self.get_text('lang_turkish'))

        lang_var = tk.StringVar(value=current_lang_name)
        lang_combo = ttk.Combobox(lang_frame, textvariable=lang_var,
                                   values=[self.get_text('lang_turkish'), self.get_text('lang_english'), self.get_text('lang_arabic')], state='readonly')
        lang_combo.pack(padx=10, pady=5)

        mode_frame = ttk.LabelFrame(general_frame, text=self.get_text('capture_mode_settings'))
        mode_frame.pack(fill='x', pady=5)
        mode_var = tk.StringVar(value=("gif" if self.config.get("gif_capture") else "ocr" if self.config.get("auto_ocr") else "screenshot"))
        modes = [
            (self.get_text('ocr_mode'), "ocr"),
            (self.get_text('gif_mode'), "gif"),
            (self.get_text('screenshot_mode'), "screenshot")
        ]
        for label, val in modes:
            ttk.Radiobutton(mode_frame, text=label, variable=mode_var, value=val).pack(anchor='w', padx=10)

        sound_frame = ttk.LabelFrame(general_frame, text=self.get_text('sound_settings'))
        sound_frame.pack(fill='x', pady=5)
        click_sound_var = tk.BooleanVar(value=self.config.get("click_sound", True))
        ttk.Checkbutton(sound_frame, text=self.get_text('click_sound'), variable=click_sound_var).pack(anchor='w', padx=10)

        # İmleç ve Kısayol sekmesi
        input_frame = ttk.Frame(notebook)
        notebook.add(input_frame, text=self.get_text('settings_input_cursor'))

        cursor_frame = ttk.LabelFrame(input_frame, text=self.get_text('mouse_cursor'))
        cursor_frame.pack(fill='x', pady=5)
        show_cursor_var = tk.BooleanVar(value=self.config.get("show_cursor", False))
        ttk.Checkbutton(cursor_frame, text=self.get_text('show_cursor'), variable=show_cursor_var).pack(anchor='w', padx=10)
        cursor_trail_var = tk.BooleanVar(value=self.config.get("cursor_trail", False))
        ttk.Checkbutton(cursor_frame, text=self.get_text('cursor_trail'), variable=cursor_trail_var).pack(anchor='w', padx=10)

        hotkey_frame = ttk.LabelFrame(input_frame, text=self.get_text('hotkeys'))
        hotkey_frame.pack(fill='x', pady=5)
        ttk.Label(hotkey_frame, text=self.get_text('capture_hotkey')).pack(anchor='w', padx=10, pady=(5,0))
        capture_hotkey_var = tk.StringVar(value=self.config.get("capture_hotkey", "ctrl+shift+f"))
        ttk.Entry(hotkey_frame, textvariable=capture_hotkey_var, width=30).pack(padx=10, pady=2)
        ttk.Label(hotkey_frame, text=self.get_text('esc_note'), 
                 font=('Arial', 9, 'italic')).pack(anchor='w', padx=10, pady=(5,0))

        # Gelişmiş sekmesi
        advanced_frame = ttk.Frame(notebook)
        notebook.add(advanced_frame, text=self.get_text('settings_advanced'))

        tray_frame = ttk.LabelFrame(advanced_frame, text=self.get_text('system_tray'))
        tray_frame.pack(fill='x', pady=5)
        minimize_to_tray_var = tk.BooleanVar(value=self.config.get("minimize_to_tray", True))
        ttk.Checkbutton(tray_frame, text=self.get_text('minimize_to_tray'), variable=minimize_to_tray_var).pack(anchor='w', padx=10)

        db_frame = ttk.LabelFrame(advanced_frame, text=self.get_text('data_management'))
        db_frame.pack(fill='x', pady=5)
        def clear_database():
            if messagebox.askyesno(self.get_text('confirm_clear_all').split('.')[0], self.get_text('confirm_clear_all')):
                self.db.clear_all_records()
                messagebox.showinfo("Success", self.get_text('records_cleared'))
        ttk.Button(db_frame, text=self.get_text('clear_all_records'), command=clear_database).pack(padx=10, pady=5)

        def save_settings():
            # UI dil değişikliği kontrolü
            selected_ui_lang = ui_lang_var.get()
            lang_map = {'Türkçe': 'tr', 'English': 'en', 'العربية': 'ar'}
            new_ui_lang = lang_map.get(selected_ui_lang, 'tr')
            
            ui_language_changed = new_ui_lang != self.current_language

            # Dil adını Tesseract koduna çevir
            selected_lang_name = lang_var.get()

            # Mevcut dil bazında kod haritası
            if self.current_language == 'tr':
                name_to_code = {'Türkçe': 'tur', 'İngilizce': 'eng', 'Arapça': 'ara'}
            elif self.current_language == 'en':
                name_to_code = {'Turkish': 'tur', 'English': 'eng', 'Arabic': 'ara'}
            else:  # Arabic
                name_to_code = {'التركية': 'tur', 'الإنجليزية': 'eng', 'العربية': 'ara'}

            self.config["language"] = name_to_code.get(selected_lang_name, 'tur')
            self.config["ui_language"] = new_ui_lang
            self.config["click_sound"] = click_sound_var.get()
            self.config["show_cursor"] = show_cursor_var.get()
            self.config["cursor_trail"] = cursor_trail_var.get()
            self.config["capture_hotkey"] = capture_hotkey_var.get()
            self.config["minimize_to_tray"] = minimize_to_tray_var.get()
            self.config["auto_ocr"] = (mode_var.get() == "ocr")
            self.config["gif_capture"] = (mode_var.get() == "gif")

            config_path = os.path.expanduser(os.path.join(
                self.config['base_directory'], self.config['directories']['config'], "config.json"))
            with open(config_path, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=4, ensure_ascii=False)

            # ComboBox'ı da güncelle
            selected_lang_name = lang_var.get()
            self.lang_combo.set(selected_lang_name)
            self.capture_mode.set(mode_var.get())

            # UI dili değiştiyse ana penceredeki combo'yu da güncelle
            if ui_language_changed:
                self.current_language = new_ui_lang
                self.ui_lang_combo.set(selected_ui_lang)

            # Yardım metnini güncelle
            if hasattr(self, 'help_label'):
                current_hotkey = capture_hotkey_var.get().upper()
                help_text = self.get_text('help_hotkey', hotkey=current_hotkey)
                self.help_label.config(text=help_text)

            win.destroy()
            
            # UI dili değiştiyse arayüzü yenile
            if ui_language_changed:
                self.refresh_ui()
                self.status_var.set(self.get_text('settings_saved'))
            else:
                self.status_var.set(self.get_text('settings_saved'))

        ttk.Button(win, text=self.get_text('save_settings'), command=save_settings).pack(pady=10)

    def save_text(self, format_type):
        text = self.result_text.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning(get_text('warning'), self.get_text('no_text_to_save'))
            return

        file_types = {
            'txt': ('Text files', '*.txt'),
            'docx': ('Word files', '*.docx')
        }

        # Belgeler/fadim/documents dizinini oluştur
        save_dir = os.path.join(os.path.expanduser(self.config['base_directory']), 
                               self.config['directories']['documents'])
        os.makedirs(save_dir, exist_ok=True)

        default_name = f"text_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        initial_path = os.path.join(save_dir, default_name)

        file_path = filedialog.asksaveasfilename(
            defaultextension=f".{format_type}",
            filetypes=[file_types[format_type]],
            initialfile=default_name,
            initialdir=save_dir
        )

        if file_path:
            try:
                if format_type == 'txt':
                    with open(file_path, 'w', encoding='utf-8') as f:
                        if self.is_arabic_text(text):
                            f.write('\u202B' + text)  # RTL karakteriyle sağa yasla
                        else:
                            f.write(text)
                else:  # docx
                    doc = Document()
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(text)

                    if self.is_arabic_text(text):
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run.font.name = 'Traditional Arabic'  # veya 'Arial'
                        run.font.size = Pt(20)
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run.font.size = Pt(14)

                    doc.save(file_path)

                self.status_var.set(self.get_text('file_saved', path=file_path))
            except Exception as e:
                messagebox.showerror("Hata", f"Dosya kaydedilirken hata oluştu: {str(e)}")

    def copy_text(self):
        text = self.result_text.get(1.0, tk.END).strip()
        if text:
            self.clipboard_clear()
            self.clipboard_append(text)
            self.status_var.set(self.get_text('text_copied'))
        else:
            messagebox.showwarning(self.get_text('warning'), self.get_text('status_no_text'))

    class CreateToolTip(object):
        def __init__(self, widget, text='widget info'):
            self.widget = widget
            self.text = text
            self.widget.bind("<Enter>", self.enter)
            self.widget.bind("<Leave>", self.close)

        def enter(self, event=None):
            x, y, _, _ = self.widget.bbox("insert")
            x += self.widget.winfo_rootx() + 25
            y += self.widget.winfo_rooty() + 20

            self.tw = tk.Toplevel(self.widget)
            self.tw.wm_overrideredirect(True)
            self.tw.wm_geometry(f"+{x}+{y}")

            label = ttk.Label(self.tw, text=self.text, justify='left',
                            background='#ffffff', relief='solid', borderwidth=1)
            label.pack(ipadx=1)

        def close(self, event=None):
            if hasattr(self, "tw"):
                self.tw.destroy()

    def show_records(self):
        records_window = tk.Toplevel(self)
        records_window.title(self.get_text('records_title'))
        records_window.geometry("1000x600")
        records_window.resizable(True, True)

        # Ana frame
        main_records_frame = ttk.Frame(records_window)
        main_records_frame.pack(expand=True, fill='both', padx=10, pady=10)

        # Başlık ve bilgi
        header_frame = ttk.Frame(main_records_frame)
        header_frame.pack(fill='x', pady=(0, 10))
        ttk.Label(header_frame, text=self.get_text('records_title'), font=('Arial', 14, 'bold')).pack(side='left')

        # Ana içerik frame (yatay bölünmüş)
        content_frame = ttk.Frame(main_records_frame)
        content_frame.pack(fill='both', expand=True, pady=(0, 10))

        # Sol taraf - Liste frame
        list_frame = ttk.LabelFrame(content_frame, text=self.get_text('records_list'))
        list_frame.pack(side='left', fill='both', expand=True, padx=(0, 5))

        # Sağ taraf - Önizleme frame
        preview_frame = ttk.LabelFrame(content_frame, text=self.get_text('screenshot_preview'))
        preview_frame.pack(side='right', fill='y', padx=(5, 0))
        preview_frame.configure(width=250)
        preview_frame.pack_propagate(False)

        # Treeview ile daha iyi görünüm
        columns = ('Date', 'Source', 'Language', 'Preview')
        self.records_tree = ttk.Treeview(list_frame, columns=columns, show='headings', height=12)

        # Sütun başlıkları
        self.records_tree.heading('Date', text=self.get_text('records_date'))
        self.records_tree.heading('Source', text=self.get_text('records_source'))
        self.records_tree.heading('Language', text=self.get_text('records_language'))
        self.records_tree.heading('Preview', text=self.get_text('records_preview'))

        # Sütun genişlikleri
        self.records_tree.column('Date', width=150)
        self.records_tree.column('Source', width=100)
        self.records_tree.column('Language', width=80)
        self.records_tree.column('Preview', width=350)

        # Kayıtları yükle
        try:
            records = self.db.get_all_records()
            for record in records:
                preview = record.text[:50] + "..." if len(record.text) > 50 else record.text
                preview = preview.replace('\n', ' ').replace('\r', ' ')
                date_str = record.timestamp.strftime('%Y-%m-%d %H:%M')

                self.records_tree.insert('', 'end', values=(
                    date_str,
                    record.source,
                    record.language or 'Bilinmiyor',
                    preview
                ))
        except Exception as e:
            print(f"❌ Kayıtlar yüklenirken hata: {e}")
            messagebox.showerror("Hata", f"Kayıtlar yüklenirken hata oluştu: {str(e)}")

        # Scrollbar
        records_scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.records_tree.yview)
        self.records_tree.configure(yscrollcommand=records_scrollbar.set)

        self.records_tree.pack(side='left', fill='both', expand=True)
        records_scrollbar.pack(side='right', fill='y')

        # Önizleme alanı
        self.preview_label = ttk.Label(preview_frame, text=self.get_text('preview_select'), 
                                      font=('Arial', 10), anchor='center', justify='center')
        self.preview_label.pack(expand=True, fill='both', padx=5, pady=5)

        # Buton frame
        button_records_frame = ttk.Frame(main_records_frame)
        button_records_frame.pack(fill='x', pady=(5, 0))

        def on_record_select(event):
            selection = self.records_tree.selection()
            if selection:
                try:
                    records = self.db.get_all_records()
                    item = self.records_tree.item(selection[0])
                    index = self.records_tree.index(selection[0])
                    record = records[index]
                    self.result_text.delete(1.0, tk.END)
                    self.result_text.insert(tk.END, record.text)
                    self.update_stats()
                    self.status_var.set(self.get_text('text_loaded', timestamp=record.timestamp))
                    if self.is_arabic_text(record.text):
                        self.align_text('right')

                    # Ekran görüntüsü önizlemesini güncelle
                    update_screenshot_preview(record)
                except Exception as e:
                    print(f"❌ Kayıt seçme hatası: {e}")

        def update_screenshot_preview(record):
            """Kayıt için ekran görüntüsü önizlemesini güncelle"""
            try:
                # Ekran görüntüsünün olabileceği yolları kontrol et
                screenshots_dir = os.path.join(
                    os.path.expanduser(self.config['base_directory']), 
                    self.config['directories']['screenshots']
                )

                # Zaman damgasından dosya adı oluştur
                timestamp_str = record.timestamp.strftime('%Y%m%d_%H%M%S')
                possible_files = [
                    f"capture_{timestamp_str}.png",
                    f"capture_{timestamp_str}.gif"
                ]

                screenshot_path = None
                for filename in possible_files:
                    full_path = os.path.join(screenshots_dir, filename)
                    if os.path.exists(full_path):
                        screenshot_path = full_path
                        break

                if screenshot_path:
                    # Resmi yükle ve boyutlandır
                    from PIL import Image, ImageTk
                    image = Image.open(screenshot_path)

                    # Önizleme için boyutlandır (max 200x200)
                    image.thumbnail((200, 200), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(image)

                    self.preview_label.configure(image=photo, text="")
                    self.preview_label.image = photo  # Referansı tut
                else:
                    self.preview_label.configure(image="", text=self.get_text('preview_not_found'))
                    if hasattr(self.preview_label, 'image'):
                        del self.preview_label.image

            except Exception as e:
                self.preview_label.configure(image="", text=self.get_text('preview_error'))
                if hasattr(self.preview_label, 'image'):
                    del self.preview_label.image

        def delete_selected_record():
            selection = self.records_tree.selection()
            if selection:
                if messagebox.askyesno(self.get_text('confirm'), self.get_text('confirm_delete_record')):
                    try:
                        records = self.db.get_all_records()
                        item = self.records_tree.item(selection[0])
                        index = self.records_tree.index(selection[0])
                        record = records[index]
                        self.db.delete_record(record.id)
                        self.records_tree.delete(selection[0])
                        messagebox.showinfo(self.get_text('success'), self.get_text('record_deleted'))
                    except Exception as e:
                        messagebox.showerror(self.get_text('error'), f"{self.get_text('record_delete_error')}: {str(e)}")
            else:
                messagebox.showwarning(self.get_text('warning'), self.get_text('select_record_to_delete'))

        self.records_tree.bind('<<TreeviewSelect>>', on_record_select)

        # Butonlar
        ttk.Button(button_records_frame, text=self.get_text('delete_selected'), 
                  command=delete_selected_record).pack(side='left', padx=5)
        ttk.Button(button_records_frame, text=self.get_text('close'), 
                  command=records_window.destroy).pack(side='right', padx=5)

        # İstatistik
        try:
            record_count = len(self.db.get_all_records())
            stats_label = ttk.Label(button_records_frame, text=self.get_text('total_records', count=record_count))
            stats_label.pack(side='left', padx=(20, 5))
        except Exception as e:
            stats_label = ttk.Label(button_records_frame, text=self.get_text('record_count_error'))
            stats_label.pack(side='left', padx=(20, 5))

    def setup_system_tray(self):
        """Sistem trayını ayarla"""
        if self.system_tray.setup_tray():
            self.system_tray.run_tray()

    def on_window_close(self):
        """Pencere kapatma eventi"""
        if self.config.get('minimize_to_tray', True) and hasattr(self, 'system_tray'):
            try:
                self.system_tray.hide_to_tray()
            except:
                self.quit_app()
        else:
            self.quit_app()

    def show_region_selector(self):
        """Bölge seçiciyi göster"""
        def region_callback(region):
            if region:
                self.selected_region = region
                x1, y1, x2, y2 = region
                self.status_var.set(self.get_text('region_selected_size', width=x2-x1, height=y2-y1))
            else:
                self.selected_region = None
                self.status_var.set(self.get_text('region_cancelled_selection'))

        selector = RegionSelector(callback=region_callback)
        selector.show_selector()

    def manual_cleanup(self):
        """Manuel temizlik işlemi"""
        cleanup_window = tk.Toplevel(self)
        cleanup_window.title(self.get_text('cleanup_title'))
        cleanup_window.geometry("450x350")
        cleanup_window.resizable(False, False)

        # Disk kullanım bilgisi
        usage = self.file_manager.get_disk_usage()

        info_frame = ttk.LabelFrame(cleanup_window, text=self.get_text('disk_usage'))
        info_frame.pack(fill='x', padx=10, pady=10)

        ttk.Label(info_frame, text=self.get_text('total_size', size=usage['total_size_mb'])).pack(pady=2)
        ttk.Label(info_frame, text=self.get_text('file_count', count=usage['file_count'])).pack(pady=2)

        # Temizlik ayarları
        settings_frame = ttk.LabelFrame(cleanup_window, text=self.get_text('cleanup_settings'))
        settings_frame.pack(fill='x', padx=10, pady=10)

        days_var = tk.IntVar(value=self.file_manager.cleanup_settings['days_to_keep'])
        max_files_var = tk.IntVar(value=self.file_manager.cleanup_settings['max_files'])

        ttk.Label(settings_frame, text=self.get_text('days_to_keep')).pack(anchor='w')
        days_spin = ttk.Spinbox(settings_frame, from_=1, to=365, textvariable=days_var, width=10)
        days_spin.pack(anchor='w', pady=2)

        ttk.Label(settings_frame, text=self.get_text('max_files')).pack(anchor='w')
        files_spin = ttk.Spinbox(settings_frame, from_=10, to=1000, textvariable=max_files_var, width=10)
        files_spin.pack(anchor='w', pady=2)

        # Butonlar
        button_frame = ttk.Frame(cleanup_window)
        button_frame.pack(fill='x', padx=10, pady=10)

        def save_cleanup_settings():
            """Temizlik ayarlarını config.json'a kaydet"""
            self.file_manager.cleanup_settings['days_to_keep'] = days_var.get()
            self.file_manager.cleanup_settings['max_files'] = max_files_var.get()

            # Config dosyasını güncelle
            self.config['cleanup'] = self.file_manager.cleanup_settings
            config_path = os.path.expanduser(
                os.path.join(self.config['base_directory'], self.config['directories']['config'], "config.json")
            )
            try:
                with open(config_path, 'w', encoding='utf-8') as f:
                    json.dump(self.config, f, indent=4, ensure_ascii=False)
                messagebox.showinfo(self.get_text('success'), self.get_text('cleanup_settings_saved'))
            except Exception as e:
                messagebox.showerror(self.get_text('error'), self.get_text('cleanup_settings_error', error=str(e)))

        def run_cleanup():
            # Ayarları güncelle
            self.file_manager.cleanup_settings['days_to_keep'] = days_var.get()
            self.file_manager.cleanup_settings['max_files'] = max_files_var.get()

            messagebox.showinfo(self.get_text('confirm'), self.get_text('cleanup_started', days=days_var.get()))

            result = self.file_manager.cleanup_old_files(show_progress=True)

            # Güncel kullanım bilgisi
            new_usage = self.file_manager.get_disk_usage()
            saved_mb = usage['total_size_mb'] - new_usage['total_size_mb']

            messagebox.showinfo(self.get_text('success'), 
                              self.get_text('cleanup_completed', cleaned=result['cleaned'], saved=saved_mb))

        def delete_all_files():
            """Tüm dosyaları ve kayıtları sil"""
            if messagebox.askyesno(self.get_text('dangerous_operation'), 
                                 self.get_text('delete_all_warning')):
                try:
                    # Tüm dosyaları sil
                    screenshots_dir = os.path.join(
                        os.path.expanduser(self.config['base_directory']), 
                        self.config['directories']['screenshots']
                    )
                    documents_dir = os.path.join(
                        os.path.expanduser(self.config['base_directory']), 
                        self.config['directories']['documents']
                    )

                    import shutil
                    deleted_files = 0

                    # Screenshots klasörünü temizle
                    if os.path.exists(screenshots_dir):
                        for file in os.listdir(screenshots_dir):
                            file_path = os.path.join(screenshots_dir, file)
                            if os.path.isfile(file_path):
                                os.remove(file_path)
                                deleted_files += 1

                    # Documents klasörünü temizle
                    if os.path.exists(documents_dir):
                        for file in os.listdir(documents_dir):
                            file_path = os.path.join(documents_dir, file)
                            if os.path.isfile(file_path):
                                os.remove(file_path)
                                deleted_files += 1

                    # Veritabanını temizle
                    self.db.clear_all_records()

                    messagebox.showinfo(self.get_text('success'), self.get_text('all_data_deleted', count=deleted_files))
                    cleanup_window.destroy()

                except Exception as e:
                    messagebox.showerror(self.get_text('error'), self.get_text('delete_error', error=str(e)))

        ttk.Button(button_frame, text=self.get_text('save_cleanup_settings'), command=save_cleanup_settings).pack(side='left', padx=5)
        ttk.Button(button_frame, text=self.get_text('start_cleanup'), command=run_cleanup).pack(side='left', padx=5)
        ttk.Button(button_frame, text=self.get_text('delete_all'), command=delete_all_files).pack(side='right', padx=5)

    def update_status(self, message):
        """Durum çubuğunu güncelle"""
        self.status_var.set(message)

    def check_database_status(self):
        """Veritabanı durumunu kontrol et"""
        try:
            # Veritabanı dosyasının varlığını kontrol et
            db_path = self.db.db_path
            if os.path.exists(db_path):
                file_size = os.path.getsize(db_path)
                print(f"✅ Veritabanı dosyası mevcut: {db_path}")
                print(f"📁 Dosya boyutu: {file_size} bytes")

                # Kayıt sayısını kontrol et
                record_count = len(self.db.get_all_records())
                print(f"📊 Toplam kayıt sayısı: {record_count}")

                return True
            else:
                print(f"❌ Veritabanı dosyası bulunamadı: {db_path}")
                return False

        except Exception as e:
            print(f"❌ Veritabanı durum kontrolü hatası: {e}")
            return False

    def check_updates_startup(self):
        """Başlangıçta güncelleme kontrolü yapar"""
        def update_check_worker():
            try:
                latest_version = check_for_updates()
                if latest_version:
                    self.after(0, lambda: show_update_dialog(latest_version, self))
                else:
                    print(f"✅ FADIM güncel (v{CURRENT_VERSION})")
            except Exception as e:
                print(f"Başlangıç güncelleme kontrolü hatası: {e}")

        # Thread'de çalıştır ki UI bloklanmasın
        update_thread = threading.Thread(target=update_check_worker, daemon=True)
        update_thread.start()

    def show_help(self):
        """Kullanım kılavuzunu göster"""
        help_window = tk.Toplevel(self)
        help_window.title(self.get_text('help_title'))
        help_window.geometry("600x500")
        help_window.resizable(True, True)

        main_help_frame = ttk.Frame(help_window)
        main_help_frame.pack(fill='both', expand=True, padx=10, pady=10)

        # Temel Kullanım sekmesi
        basic_frame = ttk.Frame(main_help_frame)

        basic_text = scrolledtext.ScrolledText(basic_frame, wrap=tk.WORD, font=('Arial', 10))
        basic_text.pack(fill='both', expand=True, padx=5, pady=5)

        # Dil bazında içerik
        # Applying the requested changes.
        if self.current_language == 'tr':
            help_content = """FADIM - Kullanım Kılavuzu

🎯 Temel Özellikler:
• Ekran görüntüsü alma ve OCR ile metin okuma
• GIF kayıt özelliği (hareketli ekran kaydı)
• Çoklu dil desteği (Türkçe, İngilizce, Arapça)
• Otomatik dosya yönetimi ve temizlik

📷 Ekran Yakalama:
1. "Ekran Alıntısı Al" butonuna tıklayın veya Ctrl+Shift+F tuşlarını kullanın
2. Fareyle yakalamak istediğiniz bölgeyi seçin
3. Seçilen bölge otomatik olarak OCR işleminden geçer

🎥 GIF Kayıt:
1. "GIF Modu" seçeneğini işaretleyin
2. Kalite, süre ve FPS ayarlarını yapın
3. "Ekran Alıntısı Al" ile kaydı başlatın
4. ESC tuşu ile erken durdurabilirsiniz

💾 Dosya Kaydetme:
• TXT Kaydet: Düz metin olarak kaydeder
• DOCX Kaydet: Word belgesi olarak kaydeder
• Arapça metinler otomatik sağa yaslanır

🔧 Ayarlar:
• Dil seçimi için üst menüden "Ayarlar > Ayarları Düzenle"
• Ses efektleri, kısayol tuşları değiştirilebilir
• Otomatik temizlik ayarları"""
        elif self.current_language == 'en':
            help_content = """FADIM - User Guide

🎯 Main Features:
• Screen capture and OCR text reading
• GIF recording feature (animated screen recording)
• Multi-language support (Turkish, English, Arabic)
• Automatic file management and cleanup

📷 Screen Capture:
1. Click "Take Screenshot" button or use Ctrl+Shift+F keys
2. Select the region you want to capture with mouse
3. Selected region is automatically processed through OCR

🎥 GIF Recording:
1. Select "GIF Mode" option
2. Set quality, duration and FPS settings
3. Start recording with "Take Screenshot"
4. You can stop early with ESC key

💾 File Saving:
• Save TXT: Saves as plain text
• Save DOCX: Saves as Word document
• Arabic texts are automatically right-aligned

🔧 Settings:
• Language selection from top menu "Settings > Edit Settings"
• Sound effects, hotkeys can be changed
• Auto cleanup settings"""
        else:  # Arabic
            help_content = """فاديم - دليل المستخدم

🎯 الميزات الرئيسية:
• التقاط الشاشة وقراءة النص بـ OCR
• ميزة تسجيل GIF (تسجيل شاشة متحرك)
• دعم متعدد اللغات (التركية، الإنجليزية، العربية)
• إدارة الملفات والتنظيف التلقائي

📷 التقاط الشاشة:
1. اضغط على زر "التقط لقطة شاشة" أو استخدم مفاتيح Ctrl+Shift+F
2. اختر المنطقة التي تريد التقاطها بالماوس
3. المنطقة المحددة تتم معالجتها تلقائياً عبر OCR

🎥 تسجيل GIF:
1. اختر خيار "وضع GIF"
2. اضبط إعدادات الجودة والمدة وFPS
3. ابدأ التسجيل بـ "التقط لقطة شاشة"
4. يمكنك الإيقاف مبكراً بمفتاح ESC

💾 حفظ الملفات:
• حفظ TXT: يحفظ كنص عادي
• حفظ DOCX: يحفظ كمستند Word
• النصوص العربية تصطف تلقائياً لليمين

🔧 الإعدادات:
• اختيار اللغة من القائمة العلوية "الإعدادات > تعديل الإعدادات"
• يمكن تغيير التأثيرات الصوتية والاختصارات
• إعدادات التنظيف التلقائي"""

        basic_text.insert(tk.END, help_content)
        basic_text.config(state='disabled')

        basic_frame.pack(fill='both', expand=True)

        # Kısayol Tuşları sekmesi  
        shortcuts_frame = ttk.Frame(main_help_frame)

        shortcuts_text = scrolledtext.ScrolledText(shortcuts_frame, wrap=tk.WORD, font=('Arial', 10))
        shortcuts_text.pack(fill='both', expand=True, padx=5, pady=5)

        # Dil bazında kısayol içeriği
        if self.current_language == 'tr':
            shortcuts_content = """⌨️ Kısayol Tuşları

📷 Ctrl+Shift+F - Ekran yakalama
🎥 Ctrl+Shift+G - GIF kayıt başlat/durdur  
📋 Ctrl+C - Metni kopyala
💾 Ctrl+S - TXT olarak kaydet
📄 Ctrl+Shift+S - DOCX olarak kaydet
🗑️ Ctrl+Delete - Metni temizle
❌ Alt+F4 - Programı kapat
⚙️ F10 - Ayarlar penceresi

🎮 GIF Kaydı Kontrolü:
• ESC - GIF kaydını durdur
• F1 - Yardım penceresi

💡 İpuçları:
• Arapça metinler otomatik sağa yaslanır
• Büyük metinler için DOCX formatını tercih edin
• GIF kayıtları otomatik olarak optimize edilir
• Veritabanı düzenli olarak temizlenir"""
        elif self.current_language == 'en':
            shortcuts_content = """⌨️ Hotkeys

📷 Ctrl+Shift+F - Screen capture
🎥 Ctrl+Shift+G - Start/stop GIF recording  
📋 Ctrl+C - Copy text
💾 Ctrl+S - Save as TXT
📄 Ctrl+Shift+S - Save as DOCX
🗑️ Ctrl+Delete - Clear text
❌ Alt+F4 - Close program
⚙️ F10 - Settings window

🎮 GIF Recording Control:
• ESC - Stop GIF recording
• F1 - Help window

💡 Tips:
• Arabic texts are automatically right-aligned
• Prefer DOCX format for large texts
• GIF recordings are automatically optimized
• Database is regularly cleaned"""
        else:  # Arabic
            shortcuts_content = """⌨️ الاختصارات

📷 Ctrl+Shift+F - التقاط الشاشة
🎥 Ctrl+Shift+G - بدء/إيقاف تسجيل GIF  
📋 Ctrl+C - نسخ النص
💾 Ctrl+S - حفظ كـ TXT
📄 Ctrl+Shift+S - حفظ كـ DOCX
🗑️ Ctrl+Delete - مسح النص
❌ Alt+F4 - إغلاق البرنامج
⚙️ F10 - نافذة الإعدادات

🎮 التحكم في تسجيل GIF:
• ESC - إيقاف تسجيل GIF
• F1 - نافذة المساعدة

💡 نصائح:
• النصوص العربية تصطف تلقائياً لليمين
• فضل صيغة DOCX للنصوص الكبيرة
• تسجيلات GIF محسنة تلقائياً
• قاعدة البيانات تنظف بانتظام"""

        shortcuts_text.insert(tk.END, shortcuts_content)
        shortcuts_text.config(state='disabled')

    def show_tesseract_help(self):
        """Tesseract yardım penceresini göster"""
        show_tesseract_installation_guide(self.current_language)

    def show_about(self):
        """Hakkında penceresini göster"""
        about_window = tk.Toplevel()
        about_window.title(self.get_text('about_title'))
        about_window.geometry("450x350")
        about_window.resizable(False, False)

        main_frame = ttk.Frame(about_window)
        main_frame.pack(fill='both', expand=True, padx=20, pady=20)

        # Logo ve başlık
        try:
            icon_path = self.get_resource_path(os.path.join('core', 'fadim_icon.png'))
            print(f"[DEBUG] Hakkında icon yolu: {icon_path}")

            if os.path.exists(icon_path):
                icon_image = Image.open(icon_path).resize((64, 64), Image.Resampling.LANCZOS)
                self.about_icon = ImageTk.PhotoImage(icon_image)
                ttk.Label(main_frame, image=self.about_icon).pack(pady=(0, 10))
                print("[DEBUG] Hakkında ikonunu yüklendi.")
            else:
                print("[DEBUG] Hakkında ikonunu bulunamadı.")
        except Exception as e:
            print(f"[HATA] Hakkında ikonunu yüklenemedi: {e}")

        ttk.Label(main_frame, text="FADIM", font=('Arial', 18, 'bold')).pack()
        
        # Dil bazında slogan göster
        if self.current_language == 'en':
            # İngilizce: sadece İngilizce slogan
            ttk.Label(main_frame, text="Fast Access Digital Image-to-Text Manager", font=('Arial', 12)).pack()
            ttk.Label(main_frame, text="Capture, Read, Save — Instantly!", font=('Arial', 10, 'italic')).pack(pady=(0, 20))
        elif self.current_language == 'ar':
            # Arapça: sadece Arapça slogan
            ttk.Label(main_frame, text="مدير النص الرقمي للوصول السريع للصور", font=('Arial', 12)).pack()
            ttk.Label(main_frame, text="التقط، اقرأ، احفظ — فوراً!", font=('Arial', 10, 'italic')).pack(pady=(0, 20))
        else:
            # Türkçe: hem İngilizce hem Türkçe
            ttk.Label(main_frame, text="Fast Access Digital Image-to-Text Manager", font=('Arial', 12)).pack()
            ttk.Label(main_frame, text="Fotoğraftan Anında Dijital Metin", font=('Arial', 12)).pack(pady=(0, 20))

        # Özellikler - dil bazında
        if self.current_language == 'en':
            features_text = """🎯 Features:
• OCR text reading from screenshots
• GIF recording capability
• Multi-language support (Turkish, English, Arabic)
• Automatic file management
• System tray integration
• Customizable hotkeys"""
        elif self.current_language == 'ar':
            features_text = """🎯 الميزات:
• قراءة النص بـ OCR من لقطات الشاشة
• إمكانية تسجيل GIF
• دعم متعدد اللغات (التركية، الإنجليزية، العربية)
• إدارة الملفات التلقائية
• تكامل شريط النظام
• اختصارات قابلة للتخصيص"""
        else:
            features_text = """🎯 Özellikler:
• OCR ile ekran görüntüsünden metin okuma
• GIF kayıt özelliği
• Çoklu dil desteği (Türkçe, İngilizce, Arapça)
• Otomatik dosya yönetimi
• Sistem tray entegrasyonu
• Özelleştirilebilir kısayol tuşları"""

        ttk.Label(main_frame, text=features_text, font=('Arial', 10), justify='left').pack(anchor='w')

        # Web sitesi
        website_label = "Website:" if self.current_language == 'en' else "الموقع:" if self.current_language == 'ar' else "Web Sitesi:"
        ttk.Label(main_frame, text=website_label, font=('Arial', 10, 'bold')).pack(anchor='w', pady=(20, 5))

        # FADIM özel sayfası
        fadim_site_label = ttk.Label(main_frame, text="muallimun.com/fadim", 
                                font=('Arial', 10, 'underline'), 
                                foreground='blue', cursor='hand2')
        fadim_site_label.pack(anchor='w')
        fadim_site_label.bind("<Button-1>", lambda e: webbrowser.open(WEBSITE_URL))

        # Ana site
        main_site_label = ttk.Label(main_frame, text="muallimun.net", 
                                font=('Arial', 10, 'underline'), 
                                foreground='blue', cursor='hand2')
        main_site_label.pack(anchor='w')
        main_site_label.bind("<Button-1>", lambda e: webbrowser.open("https://www.muallimun.net"))

        # Sürüm ve güncelleme
        version_frame = ttk.Frame(main_frame)
        version_frame.pack(anchor='w', pady=(15, 0))

        ttk.Label(version_frame, text=self.get_text('version', version=CURRENT_VERSION), font=('Arial', 9)).pack(side='left')

        update_btn = ttk.Button(version_frame, text=self.get_text('update_check_manual'), 
                               command=self.manual_update_check)
        update_btn.pack(side='left', padx=(15, 0))

        # Kapat butonu
        close_text = "❌ Close" if self.current_language == 'en' else "❌ إغلاق" if self.current_language == 'ar' else "❌ Kapat"
        ttk.Button(main_frame, text=close_text, command=about_window.destroy).pack(pady=(20, 0))

    def manual_update_check(self):
        """Manuel güncelleme kontrolü"""
        def check_worker():
            try:
                latest_version = check_for_updates()
                if latest_version:
                    self.after(0, lambda: show_update_dialog(latest_version, self))
                else:
                    self.after(0, lambda: messagebox.showinfo(
                        self.get_text('up_to_date'), 
                        self.get_text('up_to_date_message', version=CURRENT_VERSION)
                    ))
            except Exception as e:
                self.after(0, lambda: messagebox.showerror(
                    self.get_text('update_check_error'), 
                    self.get_text('update_check_failed', error=str(e))
                ))

        # Loading mesajı göster
        self.status_var.set(self.get_text('checking_updates'))

        # Thread'de kontrol et
        check_thread = threading.Thread(target=check_worker, daemon=True)
        check_thread.start()

def main():
    # Replit ortamı kontrolü
    is_replit = setup_replit_display()

    if is_replit:
        print("🌐 Replit ortamında çalışıyor - VNC penceresini açmak için Run butonuna tıklayın")
        print("📱 GUI uygulaması VNC ile görüntülenecek")

    # 🆕 Tesseract kontrolü (program kapanmadan sadece uyarı verir)
    check_tesseract_available()

    try:
        app = FADIM()
        if is_replit:
            print("✅ FADIM başarıyla başlatıldı - VNC penceresini kontrol edin")
        app.mainloop()
    except Exception as e:
        print(f"❌ Uygulama başlatma hatası: {e}")
        if is_replit:
            print("💡 VNC bağlantısı kurulamıyor olabilir, lütfen Replit VNC ayarlarını kontrol edin")

if __name__ == "__main__":
    main()